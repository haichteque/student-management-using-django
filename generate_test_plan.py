from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_plan():
    document = Document()

    # Title Page
    title = document.add_heading('Test Plan Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = document.add_paragraph('Student Management System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('\n\n\n')
    document.add_paragraph('Version: 1.0')
    document.add_paragraph('Date: December 7, 2025')
    document.add_paragraph('Prepared by: QA Team')
    
    document.add_page_break()

    # 1. Introduction
    document.add_heading('1. Introduction', level=1)
    document.add_paragraph(
        "The purpose of this Test Plan is to describe the testing strategy, scope, resources, and schedule "
        "for the testing of the Student Management System. This plan identifies the items to be tested, "
        "the features to be tested, the types of testing to be performed, and the personnel responsible "
        "for testing."
    )

    # 2. Test Items
    document.add_heading('2. Test Items', level=1)
    document.add_paragraph("The following components of the Student Management System will be tested:")
    document.add_paragraph("- User Authentication (Login/Logout)", style='List Bullet')
    document.add_paragraph("- Admin (HOD) Dashboard & Functionality", style='List Bullet')
    document.add_paragraph("- Staff Dashboard & Functionality", style='List Bullet')
    document.add_paragraph("- Student Dashboard & Functionality", style='List Bullet')
    document.add_paragraph("- Attendance Management", style='List Bullet')
    document.add_paragraph("- Course & Subject Management", style='List Bullet')

    # 3. Features to be Tested
    document.add_heading('3. Features to be Tested', level=1)
    document.add_paragraph("The following features will be validated:")
    document.add_paragraph("1. Correctness of data entry and retrieval.", style='List Number')
    document.add_paragraph("2. Access control and permission validation for different user roles.", style='List Number')
    document.add_paragraph("3. Integration between frontend forms and backend database.", style='List Number')
    document.add_paragraph("4. System stability under normal load.", style='List Number')

    # 4. Features Not to be Tested
    document.add_heading('4. Features Not to be Tested', level=1)
    document.add_paragraph("- Third-party payment gateway integration (mocked).")
    document.add_paragraph("- Performance testing under high concurrency (out of scope for this phase).")

    # 5. Approach (Test Strategy)
    document.add_heading('5. Approach', level=1)
    
    document.add_heading('5.1 Unit Testing (White-Box)', level=2)
    document.add_paragraph(
        "Unit tests will be written using Pytest. We aim for high code coverage (>70%) "
        "focusing on views and business logic."
    )
    
    document.add_heading('5.2 System/UI Testing (Black-Box)', level=2)
    document.add_paragraph(
        "End-to-end UI tests will be conducted using Cypress to simulate real user interactions "
        "across critical workflows (e.g., Add Student, Take Attendance)."
    )

    # 6. Item Pass/Fail Criteria
    document.add_heading('6. Item Pass/Fail Criteria', level=1)
    document.add_paragraph(
        "A feature is considered PASSED if:"
    )
    document.add_paragraph("- All associated unit tests pass.", style='List Bullet')
    document.add_paragraph("- All associated Cypress E2E tests pass.", style='List Bullet')
    document.add_paragraph("- No critical or high-severity bugs remain open.", style='List Bullet')

    # 7. Test Deliverables
    document.add_heading('7. Test Deliverables', level=1)
    document.add_paragraph("- Test Plan Document (This document)", style='List Bullet')
    document.add_paragraph("- Test Cases (Automated Scripts)", style='List Bullet')
    document.add_paragraph("- Test Reports (Pytest & Cypress Reports)", style='List Bullet')
    document.add_paragraph("- Defect Reports (if any)", style='List Bullet')

    # 8. Environmental Needs
    document.add_heading('8. Environmental Needs', level=1)
    document.add_paragraph("Testing will be conducted in the following environments:")
    document.add_paragraph("- Local Development Environment (Dockerized)", style='List Bullet')
    document.add_paragraph("- Staging Environment (Oracle Cloud VM, Port 8001)", style='List Bullet')
    document.add_paragraph("- CI/CD Pipeline (GitHub Actions)", style='List Bullet')

    document.save('Test_Plan.docx')
    print("Test_Plan.docx created successfully.")

if __name__ == "__main__":
    create_test_plan()
