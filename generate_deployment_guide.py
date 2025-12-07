from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_deployment_guide():
    document = Document()

    # Styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title Page
    title = document.add_heading('Deployment Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = document.add_paragraph('Student Management System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('\n\n\n')
    document.add_paragraph('Version: 1.0')
    document.add_paragraph('Date: December 7, 2025')
    document.add_paragraph('Target Audience: DevOps Engineers, Developers')
    
    document.add_page_break()

    # 1. Introduction
    document.add_heading('1. Introduction', level=1)
    document.add_paragraph(
        "This document outlines the standard operating procedure for deploying the Student Management System "
        "to its Staging and Production environments. The project utilizes a branch-based CI/CD workflow "
        "powered by GitHub Actions and Oracle Cloud Infrastructure (OCI)."
    )

    # 2. Environment Overview
    document.add_heading('2. Environment Overview', level=1)
    
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Environment'
    hdr_cells[1].text = 'Branch'
    hdr_cells[2].text = 'Details'
    
    # Staging Row
    row_cells = table.add_row().cells
    row_cells[0].text = 'Staging'
    row_cells[1].text = 'staging'
    row_cells[2].text = 'URL: http://<VM_IP>:8001\nContainer: student-app-staging\nPurpose: Testing & Verification'
    
    # Production Row
    row_cells = table.add_row().cells
    row_cells[0].text = 'Production'
    row_cells[1].text = 'main'
    row_cells[2].text = 'URL: http://<VM_IP>:8000\nContainer: student-app\nPurpose: Live User Traffic'

    document.add_paragraph('\n')

    # 3. Deployment Workflow
    document.add_heading('3. Deployment Workflow', level=1)

    # 3.1 Staging
    document.add_heading('3.1 Deploying to Staging', level=2)
    document.add_paragraph(
        "The Staging environment is an exact replica of production used for final verification. "
        "Deployments are triggered automatically when code is pushed to the 'staging' branch."
    )
    
    document.add_paragraph("Step 1: Switch to Staging Branch", style='List Number')
    p = document.add_paragraph("Ensure you are on the staging branch locally:")
    p.paragraph_format.left_indent = Inches(0.5)
    code = document.add_paragraph("git checkout staging")
    code.style = 'No Spacing'
    code.paragraph_format.left_indent = Inches(0.75)
    font = code.runs[0].font
    font.name = 'Courier New'
    font.color.rgb = RGBColor(0, 0, 128)

    document.add_paragraph("Step 2: Merge Changes", style='List Number')
    p = document.add_paragraph("Merge the feature branch you want to test:")
    p.paragraph_format.left_indent = Inches(0.5)
    code = document.add_paragraph("git merge feature-branch-name")
    code.style = 'No Spacing'
    code.paragraph_format.left_indent = Inches(0.75)
    font = code.runs[0].font
    font.name = 'Courier New'
    font.color.rgb = RGBColor(0, 0, 128)

    document.add_paragraph("Step 3: Trigger Deployment", style='List Number')
    p = document.add_paragraph("Push the changes to GitHub to start the CI/CD pipeline:")
    p.paragraph_format.left_indent = Inches(0.5)
    code = document.add_paragraph("git push origin staging")
    code.style = 'No Spacing'
    code.paragraph_format.left_indent = Inches(0.75)
    font = code.runs[0].font
    font.name = 'Courier New'
    font.color.rgb = RGBColor(0, 0, 128)

    # 3.2 Production
    document.add_heading('3.2 Deploying to Production', level=2)
    document.add_paragraph(
        "Production deployments should only occur after features have been verified on Staging. "
        "The process involves merging 'staging' into 'main'."
    )

    document.add_paragraph("Step 1: Switch to Main Branch", style='List Number')
    code = document.add_paragraph("git checkout main")
    code.style = 'No Spacing'
    code.paragraph_format.left_indent = Inches(0.75)
    font = code.runs[0].font
    font.name = 'Courier New'
    font.color.rgb = RGBColor(0, 0, 128)

    document.add_paragraph("Step 2: Merge Staging", style='List Number')
    code = document.add_paragraph("git merge staging")
    code.style = 'No Spacing'
    code.paragraph_format.left_indent = Inches(0.75)
    font = code.runs[0].font
    font.name = 'Courier New'
    font.color.rgb = RGBColor(0, 0, 128)

    document.add_paragraph("Step 3: Trigger Deployment", style='List Number')
    code = document.add_paragraph("git push origin main")
    code.style = 'No Spacing'
    code.paragraph_format.left_indent = Inches(0.75)
    font = code.runs[0].font
    font.name = 'Courier New'
    font.color.rgb = RGBColor(0, 0, 128)

    # 4. Monitoring & Verification
    document.add_heading('4. Monitoring & Verification', level=1)
    document.add_paragraph(
        "After deployment, verify the application health using the following tools:"
    )
    
    document.add_heading('4.1 New Relic (Performance)', level=2)
    document.add_paragraph(
        "Log in to one.newrelic.com and check the APM dashboard. You should see two distinct services:"
    )
    document.add_paragraph("- Student Management System (Staging)", style='List Bullet')
    document.add_paragraph("- Student Management System (Production)", style='List Bullet')
    
    document.add_heading('4.2 Manual Verification', level=2)
    document.add_paragraph(
        "Visit the application URL and perform a smoke test (Login, View Dashboard). "
        "To test error reporting, visit the debug route:"
    )
    document.add_paragraph("http://<VM_IP>:8001/sentry-debug/", style='List Bullet')

    # 5. Troubleshooting & Rollback
    document.add_heading('5. Troubleshooting & Rollback', level=1)
    
    document.add_heading('5.1 Deployment Failures', level=2)
    document.add_paragraph(
        "If the GitHub Action fails, check the 'Actions' tab logs. Common issues include:"
    )
    document.add_paragraph("- Failed Tests: Fix the code and push again.", style='List Bullet')
    document.add_paragraph("- SSH Timeout: The Oracle VM might be unreachable.", style='List Bullet')
    
    document.add_heading('5.2 Rollback Procedure', level=2)
    document.add_paragraph(
        "If a bad bug reaches production, revert the commit immediately:"
    )
    code = document.add_paragraph("git revert HEAD\ngit push origin main")
    code.style = 'No Spacing'
    code.paragraph_format.left_indent = Inches(0.5)
    font = code.runs[0].font
    font.name = 'Courier New'
    font.color.rgb = RGBColor(139, 0, 0)

    document.save('Deployment_Guide.docx')
    print("Deployment_Guide.docx created successfully.")

if __name__ == "__main__":
    create_deployment_guide()
