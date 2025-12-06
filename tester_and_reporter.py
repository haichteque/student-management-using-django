import os
import subprocess
import json
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------- Configuration ----------------
TEST_DIR = "./tests"  # Your Django test folder
JSON_REPORT = "pytest_report.json"
OUTPUT_DOCX = "Pytest_Test_Report.docx"

GREEN_LIGHT = RGBColor(198, 239, 206)
GREEN_DARK = RGBColor(155, 187, 89)

# ---------------- Helpers ----------------
def run_pytest():
    """Run pytest and generate JSON report"""
    cmd = [
        "pytest",
        TEST_DIR,
        "--json-report",
        f"--json-report-file={JSON_REPORT}"
    ]
    print("Running pytest...")
    result = subprocess.run(cmd, capture_output=True, text=True)
    print(result.stdout)
    print(result.stderr)

def parse_json_report():
    """Parse pytest JSON report"""
    if not os.path.exists(JSON_REPORT):
        print(f"No JSON report found at {JSON_REPORT}")
        return []

    with open(JSON_REPORT, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    tests = []
    for item in data.get("tests", []):
        tests.append({
            "name": item.get("nodeid", ""),
            "outcome": item.get("outcome", ""),
            "duration": item.get("duration", 0),
            "keywords": item.get("keywords", [])
        })
    return tests

def set_cell_color(cell, rgb_color):
    """Set table cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    # Use hex string directly from RGBColor
    shd.set(qn('w:fill'), str(rgb_color))  # str(rgb_color) gives 'RRGGBB'
    tcPr.append(shd)


def add_colored_table(doc, test, color):
    """Add a 3x2 table for each test"""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = "Description"
    set_cell_color(hdr_cells[0], color)
    set_cell_color(hdr_cells[1], color)

    fields = {
        "Test Name": test["name"],
        "Outcome": f'{test["outcome"].upper()} ({test["duration"]:.2f}s)',
        "Keywords": ", ".join(test["keywords"])
    }

    for key, val in fields.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = val
        set_cell_color(row_cells[0], color)
        set_cell_color(row_cells[1], color)
    
    doc.add_paragraph()  # Spacer

# ---------------- Main ----------------
def main():
    run_pytest()
    tests = parse_json_report()
    if not tests:
        print("No tests found in JSON report.")
        return

    doc = Document()
    heading = doc.add_heading("Pytest Test Results Report", 0)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph(f"Source Test Directory: {TEST_DIR}\n")

    for i, test in enumerate(tests):
        color = GREEN_LIGHT if i % 2 == 0 else GREEN_DARK
        add_colored_table(doc, test, color)

    doc.save(OUTPUT_DOCX)
    print(f"\n✅ Word report generated: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
