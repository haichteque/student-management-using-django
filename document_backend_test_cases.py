import os
import json
from google import genai
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Configuration ---
try:
    client = genai.Client()
except Exception:
    print("!!! ERROR: GEMINI_API_KEY not set or client failed to initialize.")
    exit()

CODE_DIR = "tests"
OUTPUT_DOCX = "Gemini_Generated_Test_Cases.docx"

SYSTEM_PROMPT = """
You are an expert QA engineer. Generate **all relevant backend/API test cases** for the provided Python code.
Output MUST be valid JSON ONLY, in the format:

[
  {
    "name": "Test Case Name",
    "test_scenario": "...",
    "steps": "...",
    "expected_result": "..."
  }
]

Do NOT include any other text or explanation.
If you cannot generate test cases, return an empty list: []
"""

# --- Gemini call with safe JSON parsing ---
def get_test_cases_from_gemini(file_content: str, filename: str) -> list:
    prompt = f"{SYSTEM_PROMPT}\n\n--- CODE START ---\n{file_content}\n--- CODE END ---"
    try:
        response = client.models.generate_content(
            model="gemini-2.5-pro",
            contents=prompt
        )
        text = response.text.strip()

        start = text.find("[")
        end = text.rfind("]") + 1
        if start == -1 or end == -1:
            print(f"[WARNING] No JSON array found in Gemini output for {filename}")
            return []

        json_text = text[start:end]
        test_cases = json.loads(json_text)

        if not isinstance(test_cases, list):
            raise ValueError("Expected JSON array")
        return test_cases

    except Exception as e:
        print(f"[ERROR] Gemini failed or returned invalid JSON for {filename}: {e}")
        return []

# --- Table styling helpers ---
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            for key, value in kwargs[edge].items():
                element.set(qn(f'w:{key}'), str(value))
            tcPr.append(element)

def set_cell_shading(cell, color="CCFFCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_test_cases_to_word(doc, test_cases):
    shade1 = "CCFFCC"  # light green
    shade2 = "DDFFDD"  # slightly darker green

    for tc in test_cases:
        name = tc.get("name", "Unnamed Test Case")
        heading = doc.add_heading(name, level=2)
        # Force heading color to black
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        table = doc.add_table(rows=1, cols=2)
        table.autofit = True

        # Header row
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Description"
        for cell in table.rows[0].cells:
            set_cell_shading(cell, shade1)
            set_cell_border(cell,
                            top={"sz": 6, "val": "single", "color": "000000"},
                            bottom={"sz": 6, "val": "single", "color": "000000"},
                            start={"sz": 6, "val": "single", "color": "000000"},
                            end={"sz": 6, "val": "single", "color": "000000"})

        # Test case rows with alternating shades
        row_shades = [shade2, shade1]
        row_index = 0
        for field_key, field_name in [
            ("test_scenario", "Test Scenario"),
            ("steps", "Steps"),
            ("expected_result", "Expected Result")
        ]:
            if field_key in tc:
                row_cells = table.add_row().cells
                row_cells[0].paragraphs[0].add_run(field_name).bold = True
                row_cells[1].text = tc[field_key]

                shade = row_shades[row_index % 2]
                for cell in row_cells:
                    set_cell_shading(cell, shade)
                    set_cell_border(cell,
                                    top={"sz": 6, "val": "single", "color": "000000"},
                                    bottom={"sz": 6, "val": "single", "color": "000000"},
                                    start={"sz": 6, "val": "single", "color": "000000"},
                                    end={"sz": 6, "val": "single", "color": "000000"})
                row_index += 1
        doc.add_paragraph()

def generate_word_doc_from_code(code_directory, output_filename):
    if not os.path.exists(code_directory):
        print(f"!!! ERROR: Code directory '{code_directory}' does not exist.")
        return

    doc = Document()
    title = doc.add_heading('Generated Test Cases (Gemini-2.5-Pro)', 0)
    title.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph(f"Source Code Directory: {code_directory}")
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph("---")

    total_test_cases = 0

    for root, _, files in os.walk(code_directory):
        for file in files:
            if file.endswith(".py"):
                filepath = os.path.join(root, file)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                except Exception as e:
                    print(f"!!! Could not read {filepath}: {e}")
                    continue

                test_cases = get_test_cases_from_gemini(content, file)
                if test_cases:
                    add_test_cases_to_word(doc, test_cases)
                    total_test_cases += len(test_cases)
                    print(f"  -> Processed {file}, {len(test_cases)} test cases")
                else:
                    print(f"  -> No test cases generated for {file}")

    doc.save(output_filename)
    print(f"\n✅ Done! {total_test_cases} test cases saved to {output_filename}")

if __name__ == "__main__":
    generate_word_doc_from_code(CODE_DIR, OUTPUT_DOCX)
