import subprocess
import os
from google import genai
from docx import Document
from docx.shared import RGBColor

# ----------------------------
# Gemini client
# ----------------------------
try:
    client = genai.Client()
except Exception:
    print("!!! ERROR: GEMINI_API_KEY not set or client failed to initialize.")
    exit()

OUTPUT_DOCX = "Gemini_Cypress_Test_Report.docx"
SYSTEM_PROMPT = """
You are an expert QA engineer. Analyze the Cypress test output below.
Return a plain text report with each test case:

Spec: <filename>
Test Case Number: <TC-XX>
Name: <Test Case Name>
Status: Passed or Failed

Do not include any JSON or extra explanations.
"""

# ----------------------------
# Run Cypress and capture output
# ----------------------------
def run_cypress():
    print("Running Cypress tests...")
    os.makedirs("cypress/results", exist_ok=True)
    process = subprocess.Popen(
        "npx cypress run",
        shell=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT
    )
    output, _ = process.communicate()
    return output.decode(errors="replace")

# ----------------------------
# Gemini call
# ----------------------------
def get_gemini_summary(cypress_output):
    prompt = SYSTEM_PROMPT + "\n\nCYPRESS OUTPUT:\n" + cypress_output
    try:
        response = client.models.generate_content(
            model="gemini-2.5-pro",
            contents=prompt
        )
        text = response.text.strip()
        return text
    except Exception as e:
        print(f"[ERROR] Gemini failed: {e}")
        return "Failed to get Gemini summary."

# ----------------------------
# Write to Word
# ----------------------------
def write_report_to_word(text, output_file):
    doc = Document()
    doc.add_heading("Cypress Test Report (Gemini)", 0).runs[0].font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph(text)
    doc.save(output_file)
    print(f"\n✅ Report saved to {output_file}")

# ----------------------------
# Main
# ----------------------------
if __name__ == "__main__":
    output = run_cypress()
    gemini_text = get_gemini_summary(output)
    write_report_to_word(gemini_text, OUTPUT_DOCX)
